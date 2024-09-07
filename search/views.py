import io
import json
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, FileResponse, Http404
from django.conf import settings
from django.core.files.storage import FileSystemStorage, default_storage
from django.views.decorators.csrf import csrf_exempt
from .models import File
from .utils import convert_to_txt, summarize_content_with_bart, perform_semantic_search
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from PIL import Image, ImageDraw, ImageFont
import docx
from g4f.client import Client
import os
from docx import Document
from io import StringIO, BytesIO
import base64
import logging
from .tasks import generate_summary_and_save

logger = logging.getLogger(__name__)

def search_page(request):
    users = User.objects.all()

    return render(request, 'search.html', {'users': users})


def list_files(request):
    user = request.user

    user_files = File.objects.filter(user=user)

    return render(request, 'listfiles.html', {'user_files': user_files})


# Upload new document using celery
@login_required
def insert_file1(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        fs = FileSystemStorage(location=settings.MEDIA_ROOT)
        filename = fs.save(file.name, file)
        uploaded_file_path = os.path.join(settings.MEDIA_ROOT, filename)

        text_content = convert_to_txt(uploaded_file_path)

        # Generate a preview image based on file type
        if file.name.lower().endswith('.pdf'):
            preview_url = generate_pdf_preview(uploaded_file_path)
        elif file.name.lower().endswith('.docx'):
            preview_url = generate_docx_preview(uploaded_file_path)
        else:
            preview_url = ""
        # Save file data to the database using the File model
        user = request.user
        file_instance = File.objects.create(user=user, name=file.name, path=uploaded_file_path, content=text_content, summary="", preview_url=preview_url)

        # Trigger Celery task to generate summary asynchronously
        generate_summary_and_save.delay(file_instance.id)
        logger.info(f"File upload successful and task started for file ID: {file_instance.id}")

        return redirect('/search/list_files/')
    else:
        return render(request, 'listfiles.html')


# Upload new document without using celery
@login_required
def insert_file(request):
    if request.method == 'POST' and request.FILES['file']:
        file = request.FILES['file']
        fs = FileSystemStorage(location=settings.MEDIA_ROOT)
        filename = fs.save(file.name, file)
        uploaded_file_path = os.path.join(settings.MEDIA_ROOT, filename)
        text_content = convert_to_txt(uploaded_file_path)
        if text_content is not None:

            summary = summarize_content_with_bart(text_content)

            # Generate a preview image based on file type
            if file.name.lower().endswith('.pdf'):
                preview_url = generate_pdf_preview(uploaded_file_path)
            elif file.name.lower().endswith('.docx'):
                preview_url = generate_docx_preview(uploaded_file_path)
            else:
                preview_url = ""

            user = request.user
            # Save file data to the database using the File model
            file_instance = File.insert_file_to_db(user=user, name=file.name, path=uploaded_file_path, content=text_content,
                                                   summary=summary, preview_url=preview_url)
            return redirect('/search/list_files/')
        else:
            return JsonResponse({'error': 'Unable to retrieve text content from the file.'}, status=500)
    else:
        return redirect('/search/list_files/')


def generate_pdf_preview(file_path):
    from pdf2image import convert_from_path
    from django.core.files.storage import default_storage

    poppler_path = r'C:\\poppler-24.02.0\\Library\\bin'
    images = convert_from_path(file_path, first_page=1, last_page=1, poppler_path=poppler_path)
    if images:
        preview_image = images[0]
        preview_image_path = f'previews/{os.path.basename(file_path)}.png'
        preview_image.save(os.path.join(settings.MEDIA_ROOT, preview_image_path))
        return default_storage.url(preview_image_path)
    return ""


def generate_docx_preview(file_path):
    doc = docx.Document(file_path)
    preview_text = ""
    images = []

    # Extract text and images from the document
    for para in doc.paragraphs[:10]:
        preview_text += para.text + "\n"

    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            images.append(doc.part.rels[rel].target_part.blob)

    # Create an image with PIL
    preview_image_width = 800
    preview_image_height = 1000
    image = Image.new('RGB', (preview_image_width, preview_image_height), color=(255, 255, 255))
    draw = ImageDraw.Draw(image)

    # Load a font that supports Unicode characters
    font_path = os.path.join(settings.BASE_DIR, 'C:\\dejavu-sans\\DejaVuSans.ttf')  # Update this path
    font = ImageFont.truetype(font_path, size=14)

    # Draw the preview text on the image
    current_height = 10
    margin = 10

    for line in preview_text.split('\n'):
        bbox = draw.textbbox((margin, current_height), line, font=font)
        text_height = bbox[3] - bbox[1]
        draw.text((margin, current_height), line, fill=(0, 0, 0), font=font)
        current_height += text_height + margin
        if current_height >= preview_image_height - margin:
            break

    # Add images to the preview
    for img_data in images:
        img = Image.open(io.BytesIO(img_data))
        img.thumbnail((preview_image_width - 2 * margin, preview_image_height - current_height - margin))
        image.paste(img, (margin, current_height))
        current_height += img.height + margin
        if current_height >= preview_image_height - margin:
            break

    # Save the preview image
    preview_image_path = f'previews/{os.path.basename(file_path)}.png'
    preview_image_full_path = os.path.join(settings.MEDIA_ROOT, preview_image_path)
    image.save(preview_image_full_path)

    return default_storage.url(preview_image_path)


@csrf_exempt
def semantic_search(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            query = data.get('query')
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON data in request body.'}, status=400)

        if not query:
            return JsonResponse({'error': 'Query parameter is required!'}, status=400)

        # Perform semantic search to get search results
        search_results = perform_semantic_search(query)

        if search_results:

            #top_document = search_results[0][2]
            top_document = [(result[1], result[2]) for result in search_results[:3]]

            try:
                client = Client()
                response = client.chat.completions.create(

                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant. Please respond in English."},
                        {"role": "user", "content": f"The following document has the answer to the query '{query}': {top_document}. Please provide a response in english to the query based on this document."}
                    ],
                    max_tokens=150
                )
                generated_response = response.choices[0].message.content
            except Exception as e:
                return JsonResponse({'error': f"OpenAI API error: {e}"}, status=500)

            return JsonResponse({
                'generated_response': generated_response,
                'search_results': search_results
            }, safe=False)
        else:
            return JsonResponse({'error': 'No search results found.'}, status=404)
    else:
        return JsonResponse({'error': 'Invalid request method.'}, status=405)


def update_file(request, file_id):
    if request.method == 'POST':
        new_file_name = request.POST.get('new_file_name')
        file = get_object_or_404(File, pk=file_id)
        file.name = new_file_name
        file.save()
        return render(request, 'listfiles.html')
    else:
        return JsonResponse({'error': 'Invalid request method.'}, status=405)


def delete_file(request, file_id):
    if request.method == 'POST':
        file = get_object_or_404(File, pk=file_id)
        file.delete()
        return render(request, 'listfiles.html')
    else:
        return JsonResponse({'error': 'Invalid request method.'}, status=405)


def docx_to_html(docx_path):
    document = Document(docx_path)
    html_content = StringIO()

    for para in document.paragraphs:
        html_content.write(f"<p>{para.text}</p>")

    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob
            image = Image.open(BytesIO(image_data))
            buffered = BytesIO()
            image.save(buffered, format="PNG")
            img_str = base64.b64encode(buffered.getvalue()).decode()
            img_tag = f'<img src="data:image/png;base64,{img_str}" />'
            html_content.write(img_tag)

    return html_content.getvalue()


def serve_document(request, file_name):
    file_path = os.path.join('media', file_name)
    if not os.path.exists(file_path):
        raise Http404("File does not exist")

    if file_name.endswith('.docx'):
        html_content = docx_to_html(file_path)
        return render(request, 'display_docx.html', {'html_content': html_content})

    raise Http404("File is not a .docx file")


def classify_intent(query):
    query = query.strip().lower()
    greetings = ["hi", "hello", "hey"]
    more_details_phrases = ["more details", "explain more", "i didn't understand", "can you clarify", "tell me more", "what do you mean"]

    if any(greet in query for greet in greetings):
        return "greeting"
    if any(phrase in query for phrase in more_details_phrases):
        return "more_details"

    return "general_query"


# Not the finished version, just a test of the data
client = Client()
@csrf_exempt
def chatbot(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            query = data.get('query', '').strip().lower()
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON data in request body.'}, status=400)

        if not query:
            return JsonResponse({'error': 'Query parameter is required!'}, status=400)

        # Classify the intent of the query
        intent = classify_intent(query)

        # Handle different intents
        if intent == "greeting":
            return JsonResponse({'generated_response': "Hi! How can I help you today?"}, safe=False)
        elif intent == "more_details":
            last_query = request.session.get('last_query')
            if not last_query:
                return JsonResponse({'error': 'No previous query found to provide more details.'}, status=400)
            query = last_query + " more details"

        # Perform semantic search to get search results
        search_results = perform_semantic_search(query)

        if search_results:
            top_document = [(result[1], result[2]) for result in search_results[:3]]  # Adjust as needed

            try:
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user",
                         "content": f"The following document has the answer to the query '{query}': {top_document}. Please provide a response in English to the query based on this document."}
                    ],
                    max_tokens=250
                )
                generated_response = response.choices[0].message.content
            except Exception as e:
                return JsonResponse({'error': f"OpenAI API error: {e}"}, status=500)

            # Store the current query in the session
            request.session['last_query'] = query

            return JsonResponse({'generated_response': generated_response}, safe=False)
        else:
            return JsonResponse({'error': 'No search results found.'}, status=404)
    else:
        return JsonResponse({'error': 'Invalid request method.'}, status=405)